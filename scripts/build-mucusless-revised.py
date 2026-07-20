# -*- coding: utf-8 -*-
"""Build the revised Mucusless Diet manuscript as a Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT = r"C:\Users\rajiv\OneDrive\Newsletter\2.0\Cantaloupe and the Mucusless Diet - Revised v4.docx"


def set_run_font(run, size=11, bold=False, italic=False, superscript=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True)
    p.paragraph_format.space_after = Pt(6)


def add_act(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)


def add_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc, parts):
    """parts: list of (text, kwargs) where kwargs may include bold, italic, cite."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Inches(0)
    for item in parts:
        if isinstance(item, str):
            run = p.add_run(item)
            set_run_font(run)
        else:
            text = item.get("t", "")
            run = p.add_run(text)
            set_run_font(
                run,
                bold=item.get("bold", False),
                italic=item.get("italic", False),
                superscript=item.get("sup", False),
            )
    return p


def T(text):
    return {"t": text}


def CIT(n):
    return {"t": str(n), "sup": True}


def I(text):
    return {"t": text, "italic": True}


def B(text):
    return {"t": text, "bold": True}


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, "How the Mucusless Diet Survived a Century")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("An investigation into how scientifically unsupported health systems persist.")
    set_run_font(run, size=11, italic=True)
    p.paragraph_format.space_after = Pt(18)

    # ---------- ACT I ----------
    add_act(doc, "Act I: The Claim")

    add_para(doc, [
        "A family member sent me an Instagram reel claiming that cantaloupe juice dissolves accumulated waste along the walls of the digestive tract. The juice works as a chemical cleanser, the creator said. The fruit's indigestible fiber then helps move the loosened waste toward the exit.",
        CIT(1),
    ])
    add_para(doc, [
        "The comments were just as interesting. Some people said they had experienced it themselves. Others wanted to know where to begin. A few were skeptical. Most seemed convinced they were looking at something mainstream medicine had overlooked.",
    ])
    add_para(doc, [
        "I had never heard the claim before.",
    ])
    add_para(doc, [
        "For the reel to be right, three things had to be true.",
    ])
    add_para(doc, [
        "The digestive tract would have to accumulate the kind of old waste he described.",
    ])
    add_para(doc, [
        "Cantaloupe juice would have to contain something capable of chemically loosening and dissolving that material.",
    ])
    add_para(doc, [
        "And cantaloupe fiber would have to move the loosened waste by rubbing against the digestive tract.",
    ])

    add_para(doc, [
        "The first stop was the waste itself.",
    ])
    add_para(doc, [
        "The reel moved through constipation, arterial plaque, gout, and microplastics as evidence that the body accumulates unusable waste. But those examples involve different materials, in different parts of the body, through different processes. Together, they don't establish a general layer of old food waste coating the digestive tract.",
        CIT(1),
    ])
    add_para(doc, [
        "Fecal stones are real. Doctors call them fecaliths or, when they become large masses, fecalomas. They are hardened stool associated with constipation and impaction, not proof that everyone carries old food waste along the intestinal walls.",
        CIT(4),
    ])
    add_para(doc, [
        "The reel was right about one narrow point: the digestive tract isn't clean. It isn't supposed to be. A thin mucus layer lubricates stool, shields the intestinal lining, and helps manage the microbial community living there. It's part of a healthy colon.",
        CIT(2),
        CIT(3),
    ])
    add_para(doc, [
        "It isn't permanent either. The body keeps producing, renewing, and shedding it. In a healthy colon, it doesn't build into decades-old sheets.",
        CIT(2),
    ])
    add_para(doc, [
        "The colon contains stool, but stool isn't a permanent layer of old waste attached to its wall.",
    ])

    add_para(doc, [
        "That brings us to the cantaloupe.",
    ])
    add_para(doc, [
        "If cantaloupe juice chemically dissolved accumulated waste, what did the dissolving? An enzyme? Another chemical? What study in people showed it happening?",
    ])
    add_para(doc, [
        "Neither showed up. Cantaloupe provides water, fiber, vitamin C, and carotenoids. Its water and fiber can change stool consistency and how quickly it moves. That isn't the same as chemically dissolving hardened waste.",
    ])
    add_para(doc, [
        "But the person in the video clearly believed it, and dozens of people in the comments seemed to believe it too.",
    ])
    add_para(doc, [
        "So where did this idea come from?",
    ])

    # ---------- ACT II ----------
    add_act(doc, "Act II: Following the Thread")

    add_para(doc, [
        "The reel gave me a name.",
    ])
    add_para(doc, [
        "Arnold Ehret.",
    ])
    add_para(doc, [
        "The creator said he'd practiced Ehret's Mucusless Diet Healing System for almost five years. To understand cantaloupe, he said, ",
        I('"you have to reject nutritional concepts altogether."'),
        " His profile led to a website dedicated to the lifestyle, with articles, videos, and educational materials.",
        CIT(1),
    ])
    add_para(doc, [
        "JC's Journals says its owner is affiliated with Mucus-Free Life and appears in its Transition Diet 101 course. The site says it's independently operated but directs readers to the organization's books, courses, and memberships.",
        CIT(5),
    ])
    add_para(doc, [
        "Arnold Ehret was born in Germany in 1866. He developed what he believed was a complete explanation for human health and disease, published books, operated a sanitarium, and attracted followers who treated his ideas as a more natural understanding of the body.",
        CIT(6),
    ])
    add_para(doc, [
        "What first looked like an unconventional dietary philosophy turned out to be an entirely different model of human physiology.",
    ])
    add_para(doc, [
        "According to Ehret, disease wasn't mainly about infection, genetics, or the mechanisms recognized by modern medicine. It came down to mucus building up inside the body. Food made the body cleaner or more obstructed. Protein meant something else. So did white blood cells. Even major organs had different jobs.",
        CIT(6),
    ])
    add_para(doc, [
        "People debate nutrition all the time, usually within the same understanding of physiology. Ehret was proposing another physiology. Recipes couldn't test it. The model itself had to hold up.",
    ])
    add_para(doc, [
        "Could it survive what we now know about the body?",
    ])

    # ---------- ACT III ----------
    add_act(doc, "Act III: Testing the Model")

    add_para(doc, [
        "The first surprise wasn't about food. It was about the circulatory system.",
    ])
    add_para(doc, [
        "Buried among Ehret's explanations was the claim that the lungs, not the heart, move blood through the body, while the heart functions primarily as a valve.",
        CIT(6),
    ])
    add_para(doc, [
        "I stopped reading. I thought I had misunderstood. I hadn't.",
    ])
    add_para(doc, [
        "This was no longer a disagreement about nutrition. It was a disagreement about circulation.",
    ])
    add_para(doc, [
        "The heart's ventricles generate the pressure that moves blood through the lungs and the body. The lungs exchange gases. A total artificial heart can replace the ventricles and circulate blood while the lungs remain in place.",
        CIT(7),
        CIT(8),
    ])

    add_para(doc, [
        "White blood cells weren't described as part of the body's defense against infection. They were presented as accumulated waste, another form of mucus the body was trying to eliminate.",
        CIT(6),
    ])
    add_para(doc, [
        "White blood cells develop from bone marrow and do specific immune jobs. When some types are missing, the risk of particular infections rises. They're not dietary mucus.",
        CIT(9),
    ])

    add_para(doc, [
        "The explanation always came back to mucus. Disease became mucus. Inflammation became mucus. Even the body's attempts to heal became mucus. It wasn't one idea inside the system. It supported nearly everything else.",
    ])

    add_para(doc, [
        "Ehret wasn't simply saying that people eat too much protein. He treated the whole idea of protein as a basic error and elevated fruit sugar as the body's main building material.",
        CIT(6),
    ])
    add_para(doc, [
        "Nine amino acids are essential for adults because the body can't make enough of them. They must come from food and help build tissues, enzymes, and antibodies. A well-planned plant-based diet can provide them. A fruit-dominant diet that rejects concentrated plant protein sources makes that harder.",
        CIT(10),
        CIT(11),
    ])

    add_para(doc, [
        "Vitamin B12 helps make blood cells and keeps nerves working. Plants don't naturally contain it, so people who eat no animal products need fortified foods or supplements. A deficiency can cause anemia and nerve injury, yet the body's stores can delay symptoms for years.",
        CIT(12),
    ])
    add_para(doc, [
        "The modern movement doesn't treat this as a gap to fill. In ",
        I("Spira Speaks"),
        ", Professor Spira frames B12 as a myth, questions what blood tests are measuring, and presents supplementation as part of a faulty way of thinking.",
        CIT(13),
    ])
    add_para(doc, [
        "Naming a nutrient did not invent the body's dependence on it.",
    ])

    add_para(doc, [
        "Ehret described the body as an air-gas engine powered by oxygen, minimizing food as the source of usable energy.",
        CIT(6),
        CIT(13),
    ])
    add_para(doc, [
        "Oxygen is necessary for cells to release energy, but it isn't the fuel. During fasting, the body uses energy stored as glycogen and fat, then increasingly draws on lean tissue. Oxygen permits combustion. It is not the wood.",
        CIT(14),
    ])

    add_para(doc, [
        "The physiology did not hold up. The movement survived.",
    ])
    add_para(doc, [
        "Many followers described feeling healthier, and some practices could plausibly help: eating more fruit, reducing highly processed foods, losing excess weight, or fasting in some circumstances. Those changes can be real without validating the mucus theory.",
    ])
    add_para(doc, [
        "How does a model of the body that doesn't hold up keep persuading intelligent people, generation after generation?",
    ])

    # ---------- ACT IV ----------
    add_act(doc, "Act IV: The System That Survived")

    add_para(doc, [
        "Ehret's ideas hadn't simply been preserved. They'd been translated into modern books, websites, videos, and communities. The language had changed. The presentation had changed. The confidence hadn't. The Mucusless Diet had become an ecosystem.",
    ])

    add_para(doc, [
        "One of the people who carried the system forward was Professor Spira.",
    ])
    add_para(doc, [
        "Public records identify Spira as Norman Michael Goecke, a jazz trombonist and ethnomusicologist who earned a PhD in musicology from Ohio State University in 2016. Ohio State describes him as the founder of Mucus-Free Life LLC.",
        CIT(15),
        CIT(16),
    ])
    add_para(doc, [
        "The doctorate is real. Its field is musicology, not nutrition, medicine, or physiology. The title ",
        I('"Professor"'),
        " can travel farther than the field that earned it.",
    ])
    add_para(doc, [
        "An archive preserves a document. An annotated edition tells readers how to understand it. What it explains, and what it leaves unexplained, become editorial choices. In a book giving health advice, those choices carry responsibility.",
    ])
    add_para(doc, [
        "Spira has spent years republishing, annotating, teaching, and expanding the Mucusless Diet Healing System. Mucus-Free Life sells annotated editions, courses, memberships, coaching, and menu guidance. Its lifestyle coaching program is listed at $2,997 and includes a health questionnaire, menu templates, food-preparation instruction, accountability groups, and modules that cover lemon-juice enemas, herbal bowel formulas, and ",
        I('"addiction"'),
        " to mucus-forming foods.",
        CIT(17),
        CIT(18),
    ])
    add_para(doc, [
        "Selling health education doesn't make it wrong. But books, courses, coaching, and community give a doctrine somewhere to live and spread.",
    ])
    add_para(doc, [
        "The reel creator did not invent the physiology. He compressed it. A century of doctrine became a confident social media explanation about cantaloupe.",
    ])

    add_para(doc, [
        "Innerclean.",
    ])
    add_para(doc, [
        "Ehret did not only write. He sold. Innerclean was an intestinal laxative marketed under his own name, packaged with a leaflet promoting the Mucusless Diet. It promised to remove ",
        I('"hardened feces, mucus, and other age-old"'),
        " waste from the intestines.",
        CIT(19),
    ])
    add_para(doc, [
        "The promise the reel now makes for cantaloupe was being sold in a box a century ago.",
    ])
    add_para(doc, [
        "It helped explain why so many testimonials sounded sincere.",
    ])
    add_para(doc, [
        "Historical analyses and later formulations of Innerclean have repeatedly involved stimulant and bulk-forming ingredients such as senna, buckthorn or related bark, agar, and psyllium. Those ingredients would be expected to change stool output. Senna stimulates colonic motility. Psyllium and agar hold water and add bulk. The resulting stool can look larger, more gelatinous, or more dramatic than usual.",
        CIT(19),
        CIT(20),
        CIT(21),
    ])
    add_para(doc, [
        "Followers were seeing something real. The laxative changed what left their bodies. The doctrine told them what it meant.",
    ])
    add_para(doc, [
        "Feeling better proves cleansing. Feeling worse proves cleansing. A normal test fits. So does an abnormal one. Opposite outcomes confirm the theory.",
    ])

    add_para(doc, [
        "The stakes changed when the same logic reached menstruation.",
    ])
    add_para(doc, [
        "Ehret taught that when a woman's body becomes perfectly clean through the diet, menstruation disappears, and he treated the monthly flow as impure waste. He described periods becoming less frequent and eventually stopping as the outcome of successful cleansing.",
        CIT(6),
    ])
    add_para(doc, [
        "In reality, periods can stop when the body isn't getting enough energy, after weight loss, with excessive exercise, or under stress. After other causes are ruled out, doctors call this functional hypothalamic amenorrhea. Over time, it can contribute to bone loss.",
        CIT(22),
    ])
    add_para(doc, [
        "A restrictive diet can stop menstruation, then turn that sign of physical strain into proof that the diet is working.",
    ])

    add_para(doc, [
        "Up to this point, I'd been thinking about adults choosing a restrictive system. The infant passages changed the responsibility question.",
    ])
    add_para(doc, [
        "In the same section of Ehret's book, he advises diluting cow's milk with water and sweetening it with milk sugar or honey, starting fruit juices and honey diluted in water as soon as possible, interpreting a healthy-looking infant's weight as possible waste of decayed milk, putting babies through the same cleansing process as adults, avoiding special protein foods, and suggesting that after weaning a child ",
        I('"could be raised on apples alone."'),
        CIT(6),
    ])
    add_para(doc, [
        "Today, parents are told not to give infants honey, cow's milk as a drink, or fruit juice before 12 months. Honey can cause infant botulism. Cow's milk doesn't have the right nutrient balance for an infant.",
        CIT(23),
        CIT(24),
    ])
    add_para(doc, [
        "Ehret wrote before many modern pediatric safeguards existed. Mucus-Free Life continues to market annotated editions, courses, and coaching that present the system as usable today.",
        CIT(17),
        CIT(18),
    ])

    add_para(doc, [
        "Ehret developed the theory. Modern editors decide how it reaches new readers. Today's creators decide how confidently it's compressed online. Those aren't the same responsibilities.",
    ])

    add_para(doc, [
        "Somewhere along the way, I realized I had stopped investigating a diet. I was investigating how ideas survive.",
    ])
    add_para(doc, [
        "This one survived through a chain: an intuitive explanation, experiences that felt like proof, teachers who modernized the doctrine, institutions that reinforced it, and social media that compressed it into certainty.",
    ])

    add_para(doc, [
        "When my family member sent me the reel, I thought I was investigating cantaloupe. I ended up investigating how a health system whose physiology doesn't hold up can survive for more than a century and continue persuading intelligent people.",
    ])
    add_para(doc, [
        "Understanding that process may be more useful than answering the original question.",
    ])
    add_para(doc, [
        "The next viral health claim probably won't involve mucus. When it arrives, the first question is not ",
        I('"Do I believe this?"'),
        " It's ",
        I('"How would I know if this were true?"'),
    ])

    # ---------- SOURCES ----------
    add_heading2(doc, "Sources")
    sources = [
        "JC's Journals [@jcs_journals]. \"Cantaloupe is one of the best foods on the planet...\" Instagram reel. July 17, 2026. https://www.instagram.com/p/Da5mC8iSAf9/",
        "Johansson ME, Larsson JM, Hansson GC. The two mucus layers of colon are organized by the MUC2 mucin, whereas the outer layer is a legislator of host-microbial interactions. Proc Natl Acad Sci USA. 2011;108(11):4659-4665. https://doi.org/10.1073/pnas.1006451107",
        "Song C, Chai Z, Chen S, Zhang H, Zhang X, Zhou Y. Intestinal mucus components and secretion mechanisms: what we do and do not know. Exp Mol Med. 2023;55:681-691. https://doi.org/10.1038/s12276-023-00960-y",
        "Cleveland Clinic. Fecalith & Fecaloma: What They Are, Symptoms and Treatment. https://my.clevelandclinic.org/health/diseases/fecalith",
        "JC's Journals. Terms of Use. Affiliation disclosure regarding Mucus-Free Life and Transition Diet 101. https://jcsjournals.com/",
        "Ehret A. Mucusless Diet Healing System. Historical text (archive.org). Biography and distinctive claims used in this article, including lungs as blood pump, mucus as disease foundation, menstruation as purification, and infant feeding guidance involving diluted cow's milk, honey, fruit juices, and apples. https://archive.org/details/profarnoldehrets00ehre",
        "Cleveland Clinic. How an Artificial Heart Works. https://my.clevelandclinic.org/health/procedures/22173-total-artificial-heart",
        "Copeland JG, et al. Cardiac replacement with a total artificial heart as a bridge to transplantation. N Engl J Med. 2004;351(9):859-867. https://doi.org/10.1056/NEJMoa040186",
        "Murphy K, Weaver C. Janeway's Immunobiology. 9th ed. Garland Science; 2016. Chapters on innate and adaptive leukocyte lineages and functions.",
        "National Research Council. Protein and Amino Acids. In: Recommended Dietary Allowances. 10th ed. National Academies Press; 1989. https://www.ncbi.nlm.nih.gov/books/NBK234922/",
        "EFSA Panel on Dietetic Products, Nutrition and Allergies. Scientific Opinion on Dietary Reference Values for protein. EFSA Journal. 2012. https://www.efsa.europa.eu/",
        "National Institutes of Health, Office of Dietary Supplements. Vitamin B12: Health Professional Fact Sheet. https://ods.od.nih.gov/factsheets/VitaminB12-HealthProfessional/",
        "Goecke NM (Prof. Spira). Spira Speaks: Dialogs and Essays on the Mucusless Diet Healing System. Section: Further Exploring the B-12 Myth. Mucus-Free Life LLC. https://www.mucusfreelife.com/wp-content/uploads/2016/06/Spira-Speaks-Breathair-3rd-ed.-Professor-Spira.pdf",
        "Cahill GF Jr. Fuel metabolism in starvation. Annu Rev Nutr. 2006;26:1-22. https://doi.org/10.1146/annurev.nutr.26.061505.111258",
        "Ohio State University School of Music. Musicology Careers: Norman (Michael) Goecke, PhD 2016. https://music.osu.edu/future/areas/musicology/careers",
        "Goecke NM. What Is at Stake in Jazz Education? Creative Black Music and the Twenty-First-Century Learning Environment. Doctoral dissertation, Ohio State University; 2016. https://etd.ohiolink.edu/",
        "Mucus-Free Life LLC. About / books, courses, coaching, and membership offerings. https://www.mucusfreelife.com/",
        "Mucus-Free Life. Mucus-free Diet & Lifestyle Coaching (Thinkific course page; listed price $2,997; modules include menu planning, enemas, herbology, and addiction framing). https://mucusfreelife.thinkific.com/courses/mucus-free-diet-lifestyle-coaching",
        "Notices of Judgment under the Food and Drugs Act. 19066. Misbranding of Innerclean. U.S. v. 125 Cartons of Innerclean. 1931. HathiTrust / FDA Notices of Judgment Collection. See also Cramp AJ. Laxatives: Inner-Clean. In: Nostrums and Quackery and Pseudo-Medicine. Vol. 3. AMA; 1936.",
        "Drugs.com. Senna Monograph for Professionals. https://www.drugs.com/monograph/senna.html",
        "McRorie JW Jr, McKeown NM. Understanding the physics of functional fibers in the gastrointestinal tract: an evidence-based approach to resolving enduring misconceptions about insoluble and soluble fiber. J Acad Nutr Diet. 2017;117(2):251-264. https://doi.org/10.1016/j.jand.2016.09.021",
        "Gordon CM, et al. Functional hypothalamic amenorrhea: an Endocrine Society clinical practice guideline. J Clin Endocrinol Metab. 2017;102(5):1413-1439. https://doi.org/10.1210/jc.2017-00131",
        "Centers for Disease Control and Prevention. Botulism Prevention: Do not feed honey to a child younger than 1 year old. https://www.cdc.gov/botulism/prevention/index.html",
        "Centers for Disease Control and Prevention. Foods and Drinks to Avoid or Limit: honey, cow's milk, and juice before 12 months. https://www.cdc.gov/infant-toddler-nutrition/foods-and-drinks/foods-and-drinks-to-avoid-or-limit.html",
    ]

    for i, src in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(f"{i}. {src}")
        set_run_font(run, size=9)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    run = note.add_run(
        "Note: Superscript numbers in the text correspond to the numbered sources above. "
        "This Word draft uses a plain Sources section. Collapsible website formatting can be added during the HTML build."
    )
    set_run_font(run, size=9, italic=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
